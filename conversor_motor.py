import os
import sys
import shutil
import subprocess
import zipfile
import sqlite3
import json
from typing import List, Dict

def encontrar_libreoffice():
    import os
    import sys
    import shutil

    # 1. Tenta achar no PATH do sistema operacional
    soffice = shutil.which("soffice")
    if soffice:
        return soffice

    # 2. Descobre a pasta real onde o programa está sendo executado
    if getattr(sys, 'frozen', False):
        # Se for o .exe compilado, pega a pasta onde o .exe está (a sua pasta 'dist')
        base_dir = os.path.dirname(sys.executable)
    else:
        # Se for o script rodando via terminal no VS Code, pega a pasta raiz do projeto
        base_dir = os.path.dirname(os.path.abspath(__file__))

    # 3. Procura o LibreOffice Portátil lado a lado com o programa
    caminho_portatil = os.path.join(base_dir, "LibreOfficePortable", "App", "libreoffice", "program", "soffice.exe")
    if os.path.exists(caminho_portatil):
        return caminho_portatil

    # 4. Fallback de segurança: Procura nas pastas de instalação padrão do Windows
    caminhos_padrao = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
    ]
    for cp in caminhos_padrao:
        if os.path.exists(cp):
            return cp

    return None

class ConversionError(Exception):
    """
    Exceção customizada para erros ocorridos durante o processo de conversão.
    """
    pass

# 1. Dicionário Global de Mapeamento de Formatos (MAPA_FORMATOS)
MAPA_FORMATOS = {
    # IMAGENS COMUNS
    ".png": [".png", ".jpg", ".webp", ".pdf", ".bmp", ".ico"],
    ".jpeg": [".png", ".jpg", ".webp", ".pdf", ".bmp", ".ico"],
    ".jpg": [".png", ".jpg", ".webp", ".pdf", ".bmp", ".ico"],
    ".webp": [".png", ".jpg", ".webp", ".pdf", ".bmp", ".ico"],
    ".gif": [".png", ".jpg", ".webp", ".pdf", ".bmp", ".ico", ".mp4", ".gif"],
    ".bmp": [".png", ".jpg", ".webp", ".pdf", ".bmp", ".ico"],
    ".tiff": [".png", ".jpg", ".webp", ".pdf", ".bmp", ".ico"],
    ".ico": [".png", ".jpg", ".webp", ".pdf", ".bmp", ".ico"],
    ".jfif": [".png", ".jpg", ".webp", ".pdf", ".bmp", ".ico"],
    
    # IMAGENS AVANÇADAS / RAW
    ".psd": [".png", ".jpg", ".pdf"],
    ".heic": [".png", ".jpg", ".pdf"],
    ".heif": [".png", ".jpg", ".pdf"],
    ".dng": [".png", ".jpg", ".pdf"],
    ".cr2": [".png", ".jpg", ".pdf"],
    ".nef": [".png", ".jpg", ".pdf"],

    # VÍDEOS (Gerenciado localmente via executável FFmpeg)
    ".mp4": [".mp4", ".mkv", ".webm", ".avi", ".gif", ".mp3", ".wav", ".flac"],
    ".mkv": [".mp4", ".mkv", ".webm", ".avi", ".gif", ".mp3", ".wav", ".flac"],
    ".webm": [".mp4", ".mkv", ".webm", ".avi", ".gif", ".mp3", ".wav", ".flac"],
    ".avi": [".mp4", ".mkv", ".webm", ".avi", ".gif", ".mp3", ".wav", ".flac"],
    ".mov": [".mp4", ".mkv", ".webm", ".avi", ".gif", ".mp3", ".wav", ".flac"],
    ".wmv": [".mp4", ".mkv", ".webm", ".avi", ".gif", ".mp3", ".wav", ".flac"],
    ".ts": [".mp4", ".mkv", ".webm", ".avi", ".gif", ".mp3", ".wav", ".flac"],
    ".mts": [".mp4", ".mkv", ".webm", ".avi", ".gif", ".mp3", ".wav", ".flac"],
    ".m2ts": [".mp4", ".mkv", ".webm", ".avi", ".gif", ".mp3", ".wav", ".flac"],
    ".mpg": [".mp4", ".mkv", ".webm", ".avi", ".gif", ".mp3", ".wav", ".flac"],
    ".mpeg": [".mp4", ".mkv", ".webm", ".avi", ".gif", ".mp3", ".wav", ".flac"],
    ".flv": [".mp4", ".mkv", ".webm", ".avi", ".gif", ".mp3", ".wav", ".flac"],
    ".3gp": [".mp4", ".mkv", ".webm", ".avi", ".gif", ".mp3", ".wav", ".flac"],

    # ÁUDIOS (Gerenciado localmente via executável FFmpeg)
    ".mp3": [".mp3", ".wav", ".flac", ".ogg", ".opus", ".m4a"],
    ".wav": [".mp3", ".wav", ".flac", ".ogg", ".opus", ".m4a"],
    ".flac": [".mp3", ".wav", ".flac", ".ogg", ".opus", ".m4a"],
    ".ogg": [".mp3", ".wav", ".flac", ".ogg", ".opus", ".m4a"],
    ".opus": [".mp3", ".wav", ".flac", ".ogg", ".opus", ".m4a"],
    ".aac": [".mp3", ".wav", ".flac", ".ogg", ".opus", ".m4a"],
    ".m4a": [".mp3", ".wav", ".flac", ".ogg", ".opus", ".m4a"],
    ".wma": [".mp3", ".wav", ".flac", ".ogg", ".opus", ".m4a"],

    # DADOS ESTRUTURADOS (Pandas/openpyxl/pyarrow)
    ".csv": [".xlsx", ".xls", ".json", ".parquet"],
    ".tsv": [".xlsx", ".xls", ".json", ".parquet"],
    ".xlsx": [".csv", ".xls", ".json", ".parquet"],
    ".xls": [".xlsx", ".csv", ".json", ".parquet"],
    ".parquet": [".xlsx", ".csv", ".json"],
    ".json": [".csv", ".xlsx", ".xls", ".xml", ".yaml", ".parquet"],

    # DOCUMENTOS GERAIS
    ".docx": [".pdf", ".txt"],
    ".doc": [".pdf", ".txt"],
    ".md": [".pdf", ".txt"],
    ".html": [".pdf", ".txt"],
    ".rtf": [".pdf", ".txt"],
    ".odt": [".pdf", ".txt", ".docx"],
    ".epub": [".pdf", ".txt"],
    ".pdf": [".docx"],
    ".ppt": [".pdf", ".txt", ".pptx"],
    ".pptx": [".pdf", ".txt"],

    # ARQUIVOS OPENDOCUMENT (LibreOffice / Pandas engine odf)
    ".ods": [".xlsx", ".xls", ".csv"],
    ".odp": [".pdf", ".txt", ".pptx"],

    # BANCOS DE DADOS E INFRAESTRUTURA
    ".sqlite": [".xlsx", ".csv", ".json"],
    ".db": [".xlsx", ".csv", ".json"],
    ".sql": [".json", ".csv"],
    ".xml": [".json"],
    ".yaml": [".json"],
    ".yml": [".json"],

    # ARQUIVOS COMPACTADOS
    ".zip": [".zip"],  # Destinado à extração/descompactação
    ".rar": [],        # Destinado à extração/descompactação
    ".7z": [],         # Destinado à extração/descompactação

    # DADOS GEOESPACIAIS
    ".kml": [".geojson", ".csv"],
    ".geojson": [".kml", ".csv"]
}


# Helper para converter XML parseado por ElementTree para dicionário de forma recursiva
def xml_para_dict_recursivo(element) -> Dict:
    result = {}
    for key, value in element.attrib.items():
        result[f"@{key}"] = value
    for child in element:
        child_data = xml_para_dict_recursivo(child)
        if child.tag in result:
            if not isinstance(result[child.tag], list):
                result[child.tag] = [result[child.tag]]
            result[child.tag].append(child_data)
        else:
            result[child.tag] = child_data
    text = element.text.strip() if element.text else ""
    if text:
        if result:
            result["#text"] = text
        else:
            return text
    return result if result else ""

# Helper para converter dicionário em elemento XML recursivamente
def dict_para_xml_recursivo(tag: str, d) -> subprocess.Popen:
    import xml.etree.ElementTree as ET
    elem = ET.Element(tag)
    if isinstance(d, dict):
        for k, v in d.items():
            if k.startswith('@'):
                elem.set(k[1:], str(v))
            elif k == '#text':
                elem.text = str(v)
            elif isinstance(v, list):
                for item in v:
                    elem.append(dict_para_xml_recursivo(k, item))
            else:
                elem.append(dict_para_xml_recursivo(k, v))
    else:
        elem.text = str(d)
    return elem

# 2. Função obter_saidas_permitidas(caminho_arquivo)
def obter_saidas_permitidas(caminho_arquivo: str) -> List[str]:
    """
    Extrai a extensão do arquivo e retorna a lista de formatos de saída permitidos.
    """
    ext = os.path.splitext(caminho_arquivo)[1].lower()
    return MAPA_FORMATOS.get(ext, [])

# 3. Funções de Execução das Conversões
def converter_imagem(origem: str, destino: str) -> None:
    try:
        from PIL import Image
    except ImportError:
        raise ConversionError("A biblioteca Pillow não está disponível no ambiente local.")

    try:
        ext_destino = os.path.splitext(destino)[1].lower()
        with Image.open(origem) as img:
            if ext_destino in (".jpg", ".jpeg", ".pdf"):
                if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
                    fundo = Image.new("RGB", img.size, (255, 255, 255))
                    if img.mode == "P":
                        img = img.convert("RGBA")
                    fundo.paste(img, mask=img.split()[3] if img.mode == "RGBA" else img.split()[1])
                    img = fundo
                elif img.mode != "RGB":
                    img = img.convert("RGB")

            format_name = "JPEG" if ext_destino in (".jpg", ".jpeg") else ext_destino.replace(".", "").upper()
            img.save(destino, format=format_name)
    except Exception as e:
        raise ConversionError(f"Erro no motor Pillow ao processar a imagem: {str(e)}")

def converter_midia_ffmpeg(origem: str, destino: str) -> None:
    # Apontando para a pasta bin que você vai copiar
    base_path = os.path.dirname(os.path.abspath(__file__))
    ffmpeg_exec = os.path.join(base_path, "ffmpeg_bin", "ffmpeg.exe")
    
    if not os.path.exists(ffmpeg_exec):
        raise ConversionError(f"FFmpeg não encontrado em {ffmpeg_exec}")

    try:
        # Força o uso de aspas em todos os caminhos para evitar erro de espaço no nome
        cmd = [ffmpeg_exec, "-i", origem, destino, "-y"]
        
        # Se for MP3, aplica as flags de áudio que definimos
        if destino.lower().endswith(".mp3"):
            cmd = [ffmpeg_exec, "-i", origem, "-vn", "-acodec", "libmp3lame", "-q:a", "2", destino, "-y"]

        print(f"[DEBUG] Executando FFmpeg: {' '.join(cmd)}")
        
        # Prepara argumentos para ocultar a janela do console no Windows
        run_kwargs = {'capture_output': True, 'text': True}
        if os.name == 'nt':
            run_kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW

        # Executa capturando a saída para vermos o erro real no seu terminal
        result = subprocess.run(cmd, **run_kwargs)
        
        if result.returncode != 0:
            # Aqui está o segredo: vamos ver o erro real que o FFmpeg está cuspindo
            raise ConversionError(f"Erro do FFmpeg (Código {result.returncode}): {result.stderr.strip()}")
            
    except Exception as e:
        if isinstance(e, ConversionError):
            raise e
        raise ConversionError(f"Erro ao chamar FFmpeg: {str(e)}")

def converter_dados(origem: str, destino: str, delimitador: str = ";") -> None:
    try:
        import pandas as pd
    except ImportError:
        raise ConversionError("A biblioteca pandas não está instalada no ambiente local.")

    ext_origem = os.path.splitext(origem)[1].lower()
    ext_destino = os.path.splitext(destino)[1].lower()

    try:
        # Leitura dos dados
        if ext_origem == ".csv":
            # O parâmetro on_bad_lines='skip' ignora linhas problemáticas
            # O parâmetro engine='python' é mais flexível que o padrão C
            df = pd.read_csv(origem, sep=delimitador, on_bad_lines='skip', engine='python')
        elif ext_origem == ".tsv":
            df = pd.read_csv(origem, sep="\t")
        elif ext_origem in (".xlsx", ".xls"):
            df = pd.read_excel(origem)
        elif ext_origem == ".parquet":
            df = pd.read_parquet(origem)
        elif ext_origem == ".ods":
            # Requer o pacote 'odfpy' instalado
            try:
                df = pd.read_excel(origem, engine="odf")
            except ImportError:
                raise ConversionError("A leitura de arquivos .ods requer a biblioteca 'odfpy'.")
        elif ext_origem == ".json":
            if ext_destino == ".xml":
                converter_xml_json(origem, destino)
                return
            df = pd.read_json(origem)
        elif ext_origem == ".xml":
            # Conversão Xml para Json
            converter_xml_json(origem, destino)
            return
        elif ext_origem in (".yaml", ".yml"):
            try:
                import yaml
            except ImportError:
                raise ConversionError("A leitura de arquivos YAML requer a biblioteca 'PyYAML'.")
            with open(origem, "r", encoding="utf-8") as f:
                dados_dict = yaml.safe_load(f)
            with open(destino, "w", encoding="utf-8") as f:
                json.dump(dados_dict, f, indent=4, ensure_ascii=False)
            return
        else:
            raise ConversionError(f"Extensão de dados de entrada desconhecida: {ext_origem}")

        # Escrita no formato de destino
        if ext_destino == ".csv":
            df.to_csv(destino, index=False, sep=delimitador)
        elif ext_destino in (".xlsx", ".xls"):
            df.to_excel(destino, index=False)
        elif ext_destino == ".parquet":
            df.to_parquet(destino, index=False)
        elif ext_destino == ".json":
            df.to_json(destino, orient="records", indent=4, force_ascii=False)
        elif ext_destino == ".xml":
            raise ConversionError("A geração direta de XML a partir de tabelas brutas requer modelagem estrutural. Utilize fontes JSON para XML.")
        elif ext_destino == ".yaml":
            try:
                import yaml
            except ImportError:
                raise ConversionError("A geração de arquivos YAML requer a biblioteca 'PyYAML'.")
            # Converte DataFrame para registros dicionário e salva como YAML
            dict_data = df.to_dict(orient="records")
            with open(destino, "w", encoding="utf-8") as f:
                yaml.dump(dict_data, f, default_flow_style=False, allow_unicode=True)
        else:
            raise ConversionError(f"Extensão de dados de destino desconhecida: {ext_destino}")
    except Exception as e:
        if isinstance(e, ConversionError):
            raise e
        raise ConversionError(f"Erro na conversão de dados: {str(e)}")

# Bancos de Dados Locais: SQLite/DB para Planilha (XLSX ou CSV)
def extrair_sqlite_para_planilha(origem: str, destino: str, delimitador: str = ";") -> None:
    """
    Lê todas as tabelas de um banco de dados SQLite local
    e exporta como abas em planilha XLSX ou múltiplos arquivos CSV.
    """
    try:
        import pandas as pd
    except ImportError:
        raise ConversionError("A extração de tabelas do SQLite requer a biblioteca 'pandas'.")

    try:
        conn = sqlite3.connect(origem)
        cursor = conn.cursor()
        
        # Obtém os nomes de todas as tabelas
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        tabelas = [row[0] for row in cursor.fetchall() if not row[0].startswith("sqlite_")]
        
        if not tabelas:
            conn.close()
            raise ConversionError("Nenhuma tabela de usuário encontrada no banco de dados SQLite.")

        ext_destino = os.path.splitext(destino)[1].lower()

        if ext_destino == ".xlsx":
            # Grava cada tabela como uma aba separada
            with pd.ExcelWriter(destino, engine="openpyxl") as writer:
                for tab in tabelas:
                    df = pd.read_sql_query(f'SELECT * FROM "{tab}"', conn)
                    df.to_excel(writer, sheet_name=tab[:31], index=False)  # Limite de 31 caracteres da aba Excel
        elif ext_destino == ".csv":
            # Se for apenas uma tabela, salva diretamente. Se houver mais, cria arquivos concatenados
            if len(tabelas) == 1:
                df = pd.read_sql_query(f'SELECT * FROM "{tabelas[0]}"', conn)
                df.to_csv(destino, index=False, sep=delimitador)
            else:
                base, ext = os.path.splitext(destino)
                for tab in tabelas:
                    df = pd.read_sql_query(f'SELECT * FROM "{tab}"', conn)
                    df.to_csv(f"{base}_{tab}{ext}", index=False, sep=delimitador)
        elif ext_destino == ".json":
            if len(tabelas) == 1:
                df = pd.read_sql_query(f'SELECT * FROM "{tabelas[0]}"', conn)
                df.to_json(destino, orient="records", indent=4, force_ascii=False)
            else:
                dados_completos = {}
                for tab in tabelas:
                    df = pd.read_sql_query(f'SELECT * FROM "{tab}"', conn)
                    dados_completos[tab] = df.to_dict(orient="records")
                with open(destino, "w", encoding="utf-8") as f:
                    json.dump(dados_completos, f, indent=4, ensure_ascii=False)
        else:
            conn.close()
            raise ConversionError(f"Destino para SQLite não suportado: {ext_destino}")

        conn.close()
    except Exception as e:
        if isinstance(e, ConversionError):
            raise e
        raise ConversionError(f"Falha ao processar banco SQLite: {str(e)}")

# Conversão Bidirecional entre XML e JSON
def converter_xml_json(origem: str, destino: str) -> None:
    ext_origem = os.path.splitext(origem)[1].lower()
    ext_destino = os.path.splitext(destino)[1].lower()

    try:
        import xml.etree.ElementTree as ET
        
        if ext_origem == ".xml" and ext_destino == ".json":
            # XML para JSON
            tree = ET.parse(origem)
            root = tree.getroot()
            dados_dict = {root.tag: xml_para_dict_recursivo(root)}
            with open(destino, "w", encoding="utf-8") as f:
                json.dump(dados_dict, f, indent=4, ensure_ascii=False)
        
        elif ext_origem == ".json" and ext_destino == ".xml":
            # JSON para XML
            with open(origem, "r", encoding="utf-8") as f:
                dados_json = json.load(f)
            
            # Espera-se um dicionário com uma chave raiz única representando a tag root
            if isinstance(dados_json, dict) and len(dados_json) == 1:
                root_tag = list(dados_json.keys())[0]
                root_elem = dict_para_xml_recursivo(root_tag, dados_json[root_tag])
            else:
                root_elem = dict_para_xml_recursivo("root", dados_json)
                
            tree = ET.ElementTree(root_elem)
            # Declaração XML para conformidade
            tree.write(destino, encoding="utf-8", xml_declaration=True)
        else:
            raise ConversionError(f"Conversão XML/JSON inválida de '{ext_origem}' para '{ext_destino}'.")
    except Exception as e:
        raise ConversionError(f"Erro na conversão XML <-> JSON: {str(e)}")

# 4. Utilitários para Arquivos Compactados (ZIP, RAR, 7z)
def descompactar_arquivo(origem: str, pasta_destino: str) -> None:
    """
    Extrai o conteúdo de um arquivo compactado localmente.
    Suporta nativamente ZIP. Para RAR e 7Z, faz stubs informativos de sistema.
    """
    ext = os.path.splitext(origem)[1].lower()
    if not os.path.exists(origem):
        raise ConversionError(f"Arquivo compactado não encontrado: {origem}")

    os.makedirs(pasta_destino, exist_ok=True)

    try:
        if ext == ".zip":
            with zipfile.ZipFile(origem, "r") as zip_ref:
                zip_ref.extractall(pasta_destino)
        elif ext in (".rar", ".7z"):
            # Procura por ferramentas do sistema locais (7z.exe ou UnRAR.exe no Windows/Linux)
            setez_exec = shutil.which("7z") or shutil.which("7za")
            if setez_exec:
                # 7z x origem -opasta_destino -y
                cmd = [setez_exec, "x", origem, f"-o{pasta_destino}", "-y"]
                result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                if result.returncode != 0:
                    raise ConversionError(f"Erro ao extrair via 7z: {result.stderr.strip()}")
            else:
                raise ConversionError(
                    f"A extração de arquivos '{ext}' requer o utilitário 7-Zip ('7z') instalado e no PATH offline.\n"
                    "Por favor, use arquivos .zip padrão ou instale o 7-Zip local."
                )
        else:
            raise ConversionError(f"Extensão de compactação não suportada para extração: {ext}")
    except Exception as e:
        if isinstance(e, ConversionError):
            raise e
        raise ConversionError(f"Erro na descompactação de {origem}: {str(e)}")

def compactar_arquivos(lista_origem: List[str], destino_zip: str) -> None:
    """
    Compacta uma lista de arquivos e diretórios locais em um único pacote .zip.
    """
    try:
        with zipfile.ZipFile(destino_zip, "w", zipfile.ZIP_DEFLATED) as zip_ref:
            for item in lista_origem:
                if not os.path.exists(item):
                    continue
                if os.path.isdir(item):
                    # Adiciona recursivamente os arquivos do diretório
                    for raiz, dirs, arquivos in os.walk(item):
                        for arq in arquivos:
                            caminho_completo = os.path.join(raiz, arq)
                            caminho_relativo = os.path.relpath(caminho_completo, os.path.dirname(item))
                            zip_ref.write(caminho_completo, caminho_relativo)
                else:
                    zip_ref.write(item, os.path.basename(item))
    except Exception as e:
        raise ConversionError(f"Falha ao criar o arquivo compactado ZIP: {str(e)}")

# Conversão para PDF/A (Usa LibreOffice Headless 'soffice' no sistema)
def converter_para_pdf_a(origem: str, destino: str) -> None:
    soffice_exec = encontrar_libreoffice()
    if not soffice_exec:
        raise ConversionError(
            "O LibreOffice não foi encontrado no sistema.\n"
            "Por favor, instale o LibreOffice ou coloque a versão portátil na pasta do aplicativo."
        )

    try:
        outdir = os.path.dirname(destino) or "."
        # soffice --headless --convert-to pdf:writer_pdf_Export --outdir <pasta> <arquivo>
        cmd = [soffice_exec, "--headless", "--convert-to", "pdf:writer_pdf_Export", "--outdir", outdir, origem]
        
        # Prepara argumentos para ocultar a janela do console no Windows
        run_kwargs = {'stdout': subprocess.PIPE, 'stderr': subprocess.PIPE, 'text': True, 'check': True}
        if os.name == 'nt':
            run_kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW

        result = subprocess.run(cmd, **run_kwargs)
        
        # O LibreOffice salva o pdf com o nome base do original na pasta informada
        nome_original_base = os.path.splitext(os.path.basename(origem))[0]
        pdf_gerado = os.path.join(outdir, f"{nome_original_base}.pdf")
        
        if os.path.exists(pdf_gerado) and os.path.abspath(pdf_gerado) != os.path.abspath(destino):
            if os.path.exists(destino):
                os.remove(destino)
            os.rename(pdf_gerado, destino)
    except Exception as e:
        raise ConversionError(f"Erro ao converter para PDF/A via LibreOffice: {str(e)}")

# 5. Dados Geoespaciais (KML e GeoJSON via GeoPandas)
def converter_geoespacial(origem: str, destino: str) -> None:
    """
    Converte arquivos de dados geográficos locais de forma offline.
    Requer geopandas, fiona, shapely.
    """
    try:
        import geopandas as gpd
        import fiona
    except ImportError:
        raise ConversionError(
            "A conversão geoespacial (.kml / .geojson) requer as bibliotecas 'geopandas' e 'fiona'.\n"
            "Por favor, realize a instalação offline das dependências correspondentes."
        )

    try:
        # Habilita driver KML no fiona
        fiona.drvsupport.supported_drivers['KML'] = 'rw'
        fiona.drvsupport.supported_drivers['LIBKML'] = 'rw'
    except Exception:
        pass

    try:
        ext_destino = os.path.splitext(destino)[1].lower()
        # Lê camada geográfica
        gdf = gpd.read_file(origem)

        if ext_destino == ".geojson":
            gdf.to_file(destino, driver="GeoJSON")
        elif ext_destino == ".kml":
            gdf.to_file(destino, driver="KML")
        elif ext_destino == ".csv":
            # Converte a coluna geométrica para representação em texto (WKT)
            if "geometry" in gdf.columns:
                gdf["geometry"] = gdf["geometry"].apply(lambda g: g.wkt if g else "")
            gdf.to_csv(destino, index=False)
        else:
            raise ConversionError(f"Extensão geoespacial de destino não suportada: {ext_destino}")
    except Exception as e:
        raise ConversionError(f"Erro no motor GeoPandas: {str(e)}")

def converter_documento(origem: str, destino: str) -> None:
    ext_origem = os.path.splitext(origem)[1].lower()
    ext_destino = os.path.splitext(destino)[1].lower()

    # Conversão de PDF para Word (DOCX) via pdf2docx
    if ext_origem == ".pdf" and ext_destino == ".docx":
        try:
            from pdf2docx import Converter
        except ImportError:
            raise ConversionError("A biblioteca 'pdf2docx' não está disponível no ambiente local.")
        try:
            cv = Converter(origem)
            cv.convert(destino, start=0, end=None)
            cv.close()
            return
        except Exception as e:
            raise ConversionError(f"Falha ao converter PDF para DOCX: {str(e)}")

    # Conversão de DOCX para PDF via docx2pdf com fallback para LibreOffice
    if ext_origem == ".docx" and ext_destino == ".pdf":
        try:
            from docx2pdf import convert
            # Executa a automação local do Word se disponível
            convert(origem, destino)
            return
        except Exception as e:
            # Fallback para o LibreOffice soffice
            try:
                converter_para_pdf_a(origem, destino)
                return
            except Exception as le:
                raise ConversionError(
                    "Falha ao converter DOCX para PDF de forma offline.\n"
                    "Certifique-se de que possui o Microsoft Word ou o LibreOffice instalado no sistema.\n"
                    f"Detalhes: COM/Word erro: {str(e)} | LibreOffice erro: {str(le)}"
                )

    # Caso queira PDF/A a partir de outros formatos gerais
    if ext_destino == ".pdf":
        converter_para_pdf_a(origem, destino)
        return

    if ext_origem == ".docx" and ext_destino == ".txt":
        try:
            import docx
        except ImportError:
            raise ConversionError("A biblioteca 'python-docx' não está disponível.")
        try:
            doc = docx.Document(origem)
            linhas = [p.text for p in doc.paragraphs]
            with open(destino, "w", encoding="utf-8") as f:
                f.write("\n".join(linhas))
        except Exception as e:
            raise ConversionError(f"Falha ao extrair texto do DOCX: {str(e)}")
    elif ext_origem == ".odt" and ext_destino == ".docx":
        # Converte ODT para DOCX usando LibreOffice headless
        soffice_exec = encontrar_libreoffice()
        if not soffice_exec:
            raise ConversionError("O LibreOffice não foi encontrado no sistema para a conversão ODT -> DOCX.")
        try:
            outdir = os.path.dirname(destino) or "."
            cmd = [soffice_exec, "--headless", "--convert-to", "docx", "--outdir", outdir, origem]
            
            # Prepara argumentos para ocultar a janela do console no Windows
            run_kwargs = {'check': True}
            if os.name == 'nt':
                run_kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW
                
            subprocess.run(cmd, **run_kwargs)
            gerado = os.path.join(outdir, os.path.splitext(os.path.basename(origem))[0] + ".docx")
            if os.path.exists(gerado) and os.path.abspath(gerado) != os.path.abspath(destino):
                if os.path.exists(destino):
                    os.remove(destino)
                os.rename(gerado, destino)
            return
        except Exception as e:
            raise ConversionError(f"Erro ao converter ODT para DOCX via LibreOffice: {str(e)}")
    else:
        raise ConversionError(
            f"A conversão de '{ext_origem}' para '{ext_destino}' ainda não possui suporte offline.\n"
            "Instale o LibreOffice localmente no PATH para suporte universal a documentos."
        )

# 6. Roteador de Processamento Principal
def converter_arquivo(origem: str, destino: str, delimitador: str = ";") -> None:
    if not os.path.exists(origem):
        raise ConversionError(f"O arquivo de origem não existe: {origem}")

    ext_origem = os.path.splitext(origem)[1].lower()
    ext_destino = os.path.splitext(destino)[1].lower()

    # O formato de origem '.zip' pode ser descompactado diretamente se o destino for um diretório ou se quisermos extrair
    if ext_origem in (".zip", ".rar", ".7z") and os.path.isdir(destino):
        descompactar_arquivo(origem, destino)
        return

    saidas_validas = obter_saidas_permitidas(origem)
    if ext_destino not in saidas_validas:
        raise ConversionError(
            f"Operação não permitida. A conversão de '{ext_origem}' para '{ext_destino}' não é suportada."
        )

    # Identificação de Categorias
    imagens_comuns = {".png", ".jpeg", ".jpg", ".webp", ".gif", ".bmp", ".tiff", ".ico", ".jfif"}
    imagens_raw = {".psd", ".heic", ".heif", ".dng", ".cr2", ".nef"}
    videos = {".mp4", ".mkv", ".webm", ".avi", ".mov", ".wmv", ".ts", ".mts", ".m2ts", ".mpg", ".mpeg", ".flv", ".3gp"}
    audios = {".mp3", ".wav", ".flac", ".ogg", ".opus", ".aac", ".m4a", ".wma"}
    dados = {".csv", ".tsv", ".xlsx", ".xls", ".parquet", ".json", ".xml", ".yaml", ".yml", ".ods"}
    databases = {".sqlite", ".db", ".sql"}
    documentos = {".docx", ".doc", ".md", ".html", ".rtf", ".odt", ".epub", ".odp", ".pdf", ".ppt", ".pptx"}
    geoespaciais = {".kml", ".geojson"}

    if ext_origem in imagens_comuns or ext_origem in imagens_raw:
        # Se a origem for GIF e o destino for um formato de vídeo (ex: mp4), usamos o FFmpeg
        if ext_origem == ".gif" and ext_destino in videos:
            converter_midia_ffmpeg(origem, destino)
        else:
            converter_imagem(origem, destino)
    elif ext_origem in videos or ext_origem in audios:
        converter_midia_ffmpeg(origem, destino)
    elif ext_origem in dados:
        converter_dados(origem, destino, delimitador)
    elif ext_origem in databases:
        extrair_sqlite_para_planilha(origem, destino, delimitador)
    elif ext_origem in documentos:
        converter_documento(origem, destino)
    elif ext_origem in geoespaciais:
        converter_geoespacial(origem, destino)
    else:
        raise ConversionError(f"Categoria de extensão sem tratamento associado: {ext_origem}")
# 7. Utilitários Avançados de PDF (PyMuPDF)
def juntar_pdfs(lista_caminhos: list, caminho_saida: str) -> None:
    try:
        import fitz
        print(f"[INFO] Iniciando junção de {len(lista_caminhos)} PDFs...")
        doc_saida = fitz.open()
        for i, caminho in enumerate(lista_caminhos):
            print(f"[INFO] Mesclando arquivo {i+1}: {caminho}")
            doc = fitz.open(caminho)
            doc_saida.insert_pdf(doc)
            doc.close()
        doc_saida.save(caminho_saida)
        doc_saida.close()
        print(f"[INFO] PDF combinado salvo com sucesso em: {caminho_saida}")
    except Exception as e:
        print(f"[ERROR] Falha ao juntar PDFs: {str(e)}")
        raise ConversionError(f"Falha ao juntar PDFs: {str(e)}")

def dividir_pdf(caminho_origem: str, diretorio_saida: str) -> None:
    try:
        import fitz
        print(f"[INFO] Iniciando divisão do PDF: {caminho_origem}")
        doc = fitz.open(caminho_origem)
        nome_base = os.path.splitext(os.path.basename(caminho_origem))[0]
        os.makedirs(diretorio_saida, exist_ok=True)
        for i in range(len(doc)):
            print(f"[INFO] Extraindo página {i+1} de {len(doc)}...")
            doc_pag = fitz.open()
            doc_pag.insert_pdf(doc, from_page=i, to_page=i)
            saida_pag = os.path.join(diretorio_saida, f"{nome_base}_pag_{i+1}.pdf")
            doc_pag.save(saida_pag)
            doc_pag.close()
        doc.close()
        print(f"[INFO] Divisão de PDF concluída. Arquivos em: {diretorio_saida}")
    except Exception as e:
        print(f"[ERROR] Falha ao dividir PDF: {str(e)}")
        raise ConversionError(f"Falha ao dividir PDF: {str(e)}")

def pdf_para_imagens(caminho_origem: str, diretorio_saida: str, formato: str = "png") -> None:
    try:
        import fitz
        print(f"[INFO] Convertendo PDF para imagens ({formato}): {caminho_origem}")
        doc = fitz.open(caminho_origem)
        nome_base = os.path.splitext(os.path.basename(caminho_origem))[0]
        os.makedirs(diretorio_saida, exist_ok=True)
        for i, page in enumerate(doc):
            print(f"[INFO] Renderizando página {i+1}...")
            pix = page.get_pixmap(dpi=150)
            caminho_img = os.path.join(diretorio_saida, f"{nome_base}_pag_{i+1}.{formato}")
            pix.save(caminho_img)
        doc.close()
        print(f"[INFO] Extração de imagens concluída. Arquivos em: {diretorio_saida}")
    except Exception as e:
        print(f"[ERROR] Falha ao extrair imagens do PDF: {str(e)}")
        raise ConversionError(f"Falha ao extrair imagens do PDF: {str(e)}")

def imagens_para_pdf(lista_imagens: list, caminho_saida: str) -> None:
    try:
        import fitz
        print(f"[INFO] Convertendo {len(lista_imagens)} imagens para um PDF...")
        doc_saida = fitz.open()
        for i, caminho in enumerate(lista_imagens):
            print(f"[INFO] Inserindo imagem {i+1}: {caminho}")
            img_doc = fitz.open(caminho)
            pdf_bytes = img_doc.convert_to_pdf()
            pdf_page = fitz.open("pdf", pdf_bytes)
            doc_saida.insert_pdf(pdf_page)
            img_doc.close()
            pdf_page.close()
        doc_saida.save(caminho_saida)
        doc_saida.close()
        print(f"[INFO] Imagens combinadas e salvas em: {caminho_saida}")
    except Exception as e:
        print(f"[ERROR] Falha ao converter imagens para PDF: {str(e)}")
        raise ConversionError(f"Falha ao converter imagens para PDF: {str(e)}")

def proteger_pdf(caminho_origem: str, caminho_saida: str, senha: str) -> None:
    try:
        import fitz
        print(f"[INFO] Aplicando criptografia AES-256 no PDF: {caminho_origem}")
        doc = fitz.open(caminho_origem)
        doc.save(caminho_saida, encryption=fitz.PDF_ENCRYPT_AES_256, owner_pw=senha, user_pw=senha)
        doc.close()
        print(f"[INFO] PDF protegido com sucesso em: {caminho_saida}")
    except Exception as e:
        print(f"[ERROR] Falha ao proteger PDF: {str(e)}")
        raise ConversionError(f"Falha ao proteger PDF: {str(e)}")

def desbloquear_pdf(caminho_origem: str, caminho_saida: str, senha: str) -> None:
    try:
        import fitz
        print(f"[INFO] Desbloqueando PDF: {caminho_origem}")
        doc = fitz.open(caminho_origem)
        if doc.is_encrypted:
            if doc.authenticate(senha):
                print("[INFO] Autenticação bem-sucedida! Descriptografando...")
                doc.save(caminho_saida)
                print(f"[INFO] PDF desbloqueado salvo em: {caminho_saida}")
            else:
                raise Exception("Senha incorreta!")
        else:
            print("[WARNING] O PDF original não estava protegido.")
            doc.save(caminho_saida)
        doc.close()
    except Exception as e:
        print(f"[ERROR] Falha ao desbloquear PDF: {str(e)}")
        raise ConversionError(f"Falha ao desbloquear PDF: {str(e)}")

# 8. Funções Extras de Processamento
def extrair_tabelas_pdf(origem: str, destino: str) -> None:
    try:
        import pdfplumber
        import pandas as pd
    except ImportError:
        raise ConversionError("As bibliotecas 'pdfplumber' e 'pandas' são necessárias para extrair tabelas de PDF.")

    try:
        print(f"[INFO] Extraindo tabelas de: {origem}")
        todas_tabelas = []
        with pdfplumber.open(origem) as pdf:
            for page in pdf.pages:
                tabelas_pagina = page.extract_tables()
                for tabela in tabelas_pagina:
                    if tabela:
                        df = pd.DataFrame(tabela[1:], columns=tabela[0])
                        todas_tabelas.append(df)
        
        if not todas_tabelas:
            raise ConversionError("Nenhuma tabela legível encontrada no PDF.")
            
        df_final = pd.concat(todas_tabelas, ignore_index=True)
        
        ext_destino = os.path.splitext(destino)[1].lower()
        if ext_destino == ".xlsx":
            df_final.to_excel(destino, index=False)
        elif ext_destino in [".csv", ".tsv"]:
            df_final.to_csv(destino, index=False, sep=";")
        else:
            raise ConversionError(f"Formato de destino para tabelas não suportado: {ext_destino}. Use .xlsx ou .csv.")
            
        print(f"[INFO] Tabelas extraídas e salvas em: {destino}")
    except Exception as e:
        if isinstance(e, ConversionError):
            raise e
        raise ConversionError(f"Falha ao extrair tabelas do PDF: {str(e)}")

def sanitizar_arquivo(origem: str, destino: str) -> None:
    ext_origem = os.path.splitext(origem)[1].lower()
    try:
        if ext_origem == ".pdf":
            import fitz
            print(f"[INFO] Sanitizando PDF (zerando metadados): {origem}")
            doc = fitz.open(origem)
            doc.set_metadata({})
            doc.save(destino)
            doc.close()
        elif ext_origem in [".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff"]:
            from PIL import Image
            print(f"[INFO] Sanitizando imagem (removendo EXIF): {origem}")
            with Image.open(origem) as img:
                dados_limpos = list(img.getdata())
                img_sem_exif = Image.new(img.mode, img.size)
                img_sem_exif.putdata(dados_limpos)
                img_sem_exif.save(destino)
        else:
            raise ConversionError(f"Formato não suportado para sanitização: {ext_origem}")
        print(f"[INFO] Arquivo sanitizado salvo em: {destino}")
    except Exception as e:
        if isinstance(e, ConversionError):
            raise e
        raise ConversionError(f"Falha ao sanitizar arquivo: {str(e)}")

def comprimir_arquivo(origem: str, destino: str) -> None:
    ext_origem = os.path.splitext(origem)[1].lower()
    try:
        if ext_origem == ".pdf":
            import fitz
            print(f"[INFO] Comprimindo PDF: {origem}")
            doc = fitz.open(origem)
            doc.save(destino, garbage=4, deflate=True, clean=True)
            doc.close()
        elif ext_origem in [".mp4", ".mkv", ".webm", ".avi", ".mov", ".wmv", ".ts"]:
            print(f"[INFO] Comprimindo vídeo: {origem}")
            import sys
            base_path = sys._MEIPASS if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__))
            ffmpeg_exec = os.path.join(base_path, "ffmpeg_bin", "ffmpeg.exe")
            if not os.path.exists(ffmpeg_exec):
                ffmpeg_exec = "ffmpeg"
                
            cmd = [
                ffmpeg_exec, "-i", origem,
                "-vcodec", "libx264", "-crf", "28", "-preset", "fast",
                "-y", destino
            ]
            run_kwargs = {'capture_output': True, 'text': True}
            if os.name == 'nt':
                run_kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW
                
            result = subprocess.run(cmd, **run_kwargs)
            if result.returncode != 0:
                raise ConversionError(f"Erro do FFmpeg (Código {result.returncode}): {result.stderr.strip()}")
        else:
            raise ConversionError(f"Formato não suportado para compressão: {ext_origem}")
        print(f"[INFO] Arquivo comprimido salvo em: {destino}")
    except Exception as e:
        if isinstance(e, ConversionError):
            raise e
        raise ConversionError(f"Falha ao comprimir arquivo: {str(e)}")

def ocr_documento(origem: str, destino: str) -> None:
    try:
        import pytesseract
        from PIL import Image
        import sys
        import io
        import os
        
        # Aponta para o Tesseract local que você incluiu
        base_path = sys._MEIPASS if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__))
        tesseract_cmd = os.path.join(base_path, "Tesseract-OCR", "tesseract.exe")
        
        if os.path.exists(tesseract_cmd):
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

        ext_origem = os.path.splitext(origem)[1].lower()
        ext_destino = os.path.splitext(destino)[1].lower()
        texto_extraido = ""
        
        # Função interna à prova de falhas para o pacote de idioma
        def extrair_texto(img_obj):
            try:
                # Tenta ler em Português primeiro
                return pytesseract.image_to_string(img_obj, lang="por")
            except pytesseract.TesseractError:
                print("[WARNING] Idioma 'por' não encontrado no Tesseract. Usando o padrão.")
                return pytesseract.image_to_string(img_obj)
        
        # --- A MÁGICA ACONTECE AQUI ---
        if ext_origem == ".pdf":
            import fitz
            print(f"[INFO] Fatiando PDF para leitura óptica: {origem}")
            doc = fitz.open(origem)
            for i, page in enumerate(doc):
                print(f"[INFO] Lendo texto da página {i+1}/{len(doc)}...")
                # dpi=300 garante altíssima resolução para o Tesseract não errar letras
                pix = page.get_pixmap(dpi=300)
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                texto_extraido += extrair_texto(img) + "\n\n"
            doc.close()
        elif ext_origem in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]:
            print(f"[INFO] Realizando OCR na imagem: {origem}")
            with Image.open(origem) as img:
                texto_extraido = extrair_texto(img)
        else:
            raise ConversionError(f"Formato não suportado para OCR: {ext_origem}")
            
        print(f"[INFO] Gerando arquivo de saída textual: {destino}")
        
        if ext_destino == ".docx":
            try:
                import docx
            except ImportError:
                raise ConversionError("A biblioteca 'python-docx' é necessária para salvar em DOCX.")
            
            documento = docx.Document()
            documento.add_heading('Texto Extraído via OCR', 0)
            
            # Adiciona o texto preservando parágrafos e quebras de linha
            for paragrafo in texto_extraido.split('\n'):
                if paragrafo.strip():
                    documento.add_paragraph(paragrafo.strip())
                    
            documento.save(destino)
        else:
            # Fallback seguro para bloco de notas (.txt)
            if not destino.endswith(".txt"):
                destino = os.path.splitext(destino)[0] + ".txt"
            with open(destino, "w", encoding="utf-8") as f:
                f.write(texto_extraido)
            
        print(f"[INFO] OCR concluído com sucesso. Salvo em: {destino}")
    except Exception as e:
        if isinstance(e, ConversionError):
            raise e
        raise ConversionError(f"Falha ao realizar OCR: {str(e)}")

def rotacionar_pdf(origem: str, destino: str, angulo: int) -> None:
    try:
        import fitz
        doc = fitz.open(origem)
        for page in doc:
            page.set_rotation(int(angulo))
        doc.save(destino)
        doc.close()
    except Exception as e:
        raise ConversionError(f"Erro ao rotacionar PDF: {str(e)}")

def numerar_paginas_pdf(origem: str, destino: str) -> None:
    try:
        import fitz
        doc = fitz.open(origem)
        for i, page in enumerate(doc):
            texto = f"Página {i+1} de {len(doc)}"
            ponto = fitz.Point(page.rect.width - 120, page.rect.height - 30)
            page.insert_text(ponto, texto, fontsize=10, color=(0, 0, 0))
        doc.save(destino)
        doc.close()
    except Exception as e:
        raise ConversionError(f"Erro ao numerar páginas do PDF: {str(e)}")

def remover_paginas_pdf(origem: str, destino: str, paginas_str: str) -> None:
    try:
        import fitz
        doc = fitz.open(origem)
        paginas_a_remover = set()
        total_paginas = len(doc)
        
        partes = [p.strip() for p in paginas_str.split(",") if p.strip()]
        if not partes:
            raise ConversionError("Nenhuma página informada para remoção. Insira pelo menos um número de página.")

        for parte in partes:
            try:
                if "-" in parte:
                    subpartes = parte.split("-")
                    if len(subpartes) == 2:
                        inicio = int(subpartes[0].strip())
                        fim = int(subpartes[1].strip())
                        if inicio <= 0 or fim <= 0:
                            raise ConversionError("Os números de páginas devem ser maiores que zero.")
                        if inicio > fim:
                            raise ConversionError(f"Intervalo inválido ({inicio}-{fim}). O início deve ser menor ou igual ao fim.")
                        for p_num in range(inicio, fim + 1):
                            paginas_a_remover.add(p_num - 1)
                    else:
                        raise ConversionError(f"Intervalo de páginas inválido: '{parte}'")
                else:
                    p_num = int(parte)
                    if p_num <= 0:
                        raise ConversionError("Os números de páginas devem ser maiores que zero.")
                    paginas_a_remover.add(p_num - 1)
            except ValueError:
                raise ConversionError(f"Valor inválido encontrado: '{parte}'. Certifique-se de preencher apenas números inteiros maiores que zero (Ex: 1, 3-5, 8).")
        
        indices_validos = sorted(
            [p for p in paginas_a_remover if 0 <= p < total_paginas],
            reverse=True
        )
        
        if not indices_validos:
            raise ConversionError(f"Nenhuma das páginas informadas ({paginas_str}) existe no documento. O PDF possui apenas {total_paginas} página(s).")

        if len(indices_validos) == total_paginas:
            doc.close()
            raise ConversionError("Não é permitido remover todas as páginas do PDF. O documento precisa ter pelo menos 1 página restante.")
            
        for p in indices_validos:
            doc.delete_page(p)
            
        doc.save(destino)
        doc.close()
    except Exception as e:
        if isinstance(e, ConversionError):
            raise e
        raise ConversionError(f"Erro ao remover páginas do PDF: {str(e)}")

def aplicar_marca_dagua(origem: str, destino: str, texto: str) -> None:
    if not origem.lower().endswith(".pdf"):
        raise ConversionError("A marca d'água só pode ser aplicada em arquivos PDF.")
    try:
        import fitz
        doc = fitz.open(origem)
        for page in doc:
            ponto = fitz.Point(50, page.rect.height / 2)
            page.insert_text(ponto, texto, fontsize=60, color=(0.6, 0.6, 0.6), fill_opacity=0.3, morph=(ponto, fitz.Matrix(45)))
        doc.save(destino)
        doc.close()
    except Exception as e:
        raise ConversionError(f"Erro ao aplicar marca d'água: {str(e)}")

def censurar_pdf(origem: str, destino: str, texto_alvo: str) -> None:
    try:
        import fitz
        doc = fitz.open(origem)
        for page in doc:
            areas = page.search_for(texto_alvo)
            for area in areas:
                page.add_redact_annot(area, fill=(0, 0, 0))
            if areas:
                page.apply_redactions()
        doc.save(destino)
        doc.close()
    except Exception as e:
        raise ConversionError(f"Erro ao censurar PDF: {str(e)}")

def extrair_paginas_pdf(origem: str, destino: str, paginas_str: str) -> None:
    try:
        import fitz
        doc_origem = fitz.open(origem)
        doc_destino = fitz.open()
        total_paginas = len(doc_origem)
        
        partes = [p.strip() for p in paginas_str.split(",") if p.strip()]
        if not partes:
            raise ConversionError("Nenhuma página informada para extração. Insira pelo menos um número de página.")

        for parte in partes:
            try:
                if "-" in parte:
                    subpartes = parte.split("-")
                    if len(subpartes) == 2:
                        inicio = int(subpartes[0].strip())
                        fim = int(subpartes[1].strip())
                        if inicio <= 0 or fim <= 0:
                            raise ConversionError("Os números de páginas devem ser maiores que zero.")
                        if inicio > total_paginas or fim > total_paginas:
                            raise ConversionError(f"Página fora dos limites do documento. O PDF possui apenas {total_paginas} página(s).")
                        if inicio > fim:
                            raise ConversionError(f"Intervalo inválido ({inicio}-{fim}). O início deve ser menor ou igual ao fim.")
                        doc_destino.insert_pdf(doc_origem, from_page=inicio-1, to_page=fim-1)
                    else:
                        raise ConversionError(f"Intervalo de páginas inválido: '{parte}'")
                else:
                    num = int(parte)
                    if num <= 0:
                        raise ConversionError("Os números de páginas devem ser maiores que zero.")
                    if num > total_paginas:
                        raise ConversionError(f"Página {num} fora dos limites do documento. O PDF possui apenas {total_paginas} página(s).")
                    doc_destino.insert_pdf(doc_origem, from_page=num-1, to_page=num-1)
            except ValueError:
                raise ConversionError(f"Valor inválido encontrado: '{parte}'. Certifique-se de preencher apenas números inteiros maiores que zero (Ex: 1, 3-5, 8).")
                
        doc_destino.save(destino)
        doc_destino.close()
        doc_origem.close()
    except Exception as e:
        if isinstance(e, ConversionError):
            raise e
        raise ConversionError(f"Erro ao extrair páginas do PDF: {str(e)}")

def reparar_pdf(origem: str, destino: str) -> None:
    errors = []
    # 1. Tenta usar o pikepdf
    try:
        import pikepdf
        print(f"[INFO] Tentando reparar PDF com pikepdf: {origem}")
        with pikepdf.open(origem) as pdf:
            pdf.save(destino)
            print("[INFO] PDF reparado com sucesso usando pikepdf.")
            return
    except Exception as e:
        errors.append(f"pikepdf: {str(e)}")

    # 2. Se o pikepdf falhar, tenta usar o PyMuPDF (fitz)
    try:
        import fitz
        print(f"[INFO] pikepdf falhou. Tentando reparar PDF com PyMuPDF (fitz): {origem}")
        doc = fitz.open(origem)
        doc.save(destino, clean=True, garbage=4, deflate=True)
        doc.close()
        print("[INFO] PDF reparado com sucesso usando PyMuPDF (fitz).")
        return
    except Exception as e:
        errors.append(f"PyMuPDF: {str(e)}")

    # Se ambos falharem
    raise ConversionError(f"Não foi possível reparar o PDF. Detalhes dos erros: {'; '.join(errors)}")

def assinar_pdf(origem: str, destino: str, certificado_pfx: str, senha_pfx: str) -> None:
    try:
        from pyhanko.sign import signers
        from pyhanko.pdf_utils.incremental_writer import IncrementalPdfFileWriter
        
        signer = signers.SimpleSigner.load_pkcs12(
            pfx_file=certificado_pfx,
            passphrase=senha_pfx.encode('utf-8')
        )
        
        with open(origem, 'rb') as inf:
            w = IncrementalPdfFileWriter(inf)
            with open(destino, 'wb') as outf:
                signers.sign_pdf(
                    w,
                    signers.PdfSignatureMetadata(field_name='Assinatura_A1'),
                    signer=signer,
                    output=outf
                )
    except Exception as e:
        raise ConversionError(f"Erro ao assinar PDF: {str(e)}")

def fatiar_pdf_por_tamanho(origem: str, destino_zip: str, limite_mb: float) -> None:
    import fitz
    import zipfile
    
    limite_bytes = limite_mb * 1024 * 1024
    origem_doc = fitz.open(origem)
    pdf_parts = []
    
    chunk = fitz.open()
    nome_base = os.path.splitext(os.path.basename(origem))[0]
    
    for i in range(len(origem_doc)):
        chunk.insert_pdf(origem_doc, from_page=i, to_page=i)
        
        if len(chunk.write()) > limite_bytes:
            if len(chunk) > 1:
                chunk.delete_page(-1)
                part_index = len(pdf_parts) + 1
                pdf_parts.append((f"{nome_base}_parte_{part_index}.pdf", chunk.write()))
                
                chunk.close()
                chunk = fitz.open()
                chunk.insert_pdf(origem_doc, from_page=i, to_page=i)
            else:
                part_index = len(pdf_parts) + 1
                pdf_parts.append((f"{nome_base}_parte_{part_index}.pdf", chunk.write()))
                chunk.close()
                chunk = fitz.open()
                
    if len(chunk) > 0:
        part_index = len(pdf_parts) + 1
        pdf_parts.append((f"{nome_base}_parte_{part_index}.pdf", chunk.write()))
        
    chunk.close()
    origem_doc.close()
    
    with zipfile.ZipFile(destino_zip, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for nome_parte, pdf_bytes in pdf_parts:
            zipf.writestr(nome_parte, pdf_bytes)

def juntar_docx(lista_caminhos: list, caminho_saida: str) -> None:
    try:
        from docx import Document
    except ImportError:
        raise ConversionError("A biblioteca 'python-docx' não está disponível.")
        
    try:
        if not lista_caminhos:
            raise ConversionError("Nenhum arquivo Word enviado para juntar.")
            
        doc_merged = Document(lista_caminhos[0])
        
        for i in range(1, len(lista_caminhos)):
            doc_merged.add_page_break()
            sub_doc = Document(lista_caminhos[i])
            for element in sub_doc.element.body:
                # Pula os elementos de propriedades de seção para manter o XML válido
                if element.tag.endswith('sectPr'):
                    continue
                doc_merged.element.body.append(element)
                
        doc_merged.save(caminho_saida)
        print(f"[INFO] Documentos Word combinados com sucesso em: {caminho_saida}")
    except Exception as e:
        print(f"[ERROR] Falha ao juntar arquivos DOCX: {str(e)}")
        raise ConversionError(f"Falha ao juntar arquivos Word: {str(e)}")

def renomear_em_lote(caminhos_origem: list, destino_zip: str, padrao: str) -> None:
    import os
    import zipfile
    
    with zipfile.ZipFile(destino_zip, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for i, caminho in enumerate(caminhos_origem):
            if not os.path.exists(caminho):
                continue
                
            nome_original = os.path.basename(caminho)
            nome_sem_ext, ext = os.path.splitext(nome_original)
            
            novo_nome = padrao.replace('{nome_original}', nome_original)
            novo_nome = novo_nome.replace('{nome}', nome_sem_ext)
            novo_nome = novo_nome.replace('{i}', str(i + 1))
            
            if ext and not novo_nome.lower().endswith(ext.lower()):
                novo_nome += ext
                
            zipf.write(caminho, arcname=novo_nome)

# 9. Interface de Linha de Comando (CLI) para Testes
if __name__ == "__main__":
    import argparse
    import sys

    parser = argparse.ArgumentParser(description="Conversor Offline - Motor Backend Expandido")
    parser.add_argument("--origem", required=True, help="Caminho do arquivo de origem ou diretório")
    parser.add_argument("--destino", required=True, help="Caminho do arquivo de destino ou diretório de extração")
    parser.add_argument("--delim", default=";", help="Delimitador utilizado para CSV (padrão ';')")
    parser.add_argument("--compactar", action="store_true", help="Gera um ZIP a partir dos caminhos fornecidos em --origem (separados por vírgula)")

    try:
        args = parser.parse_args()
        print("=" * 45)
        print("CONVERSOR OFFLINE - INICIANDO MOTOR")
        print("=" * 45)
        
        if args.compactar:
            caminhos_lista = [c.strip() for c in args.origem.split(",")]
            print(f"Compactando: {caminhos_lista}")
            print(f"Destino ZIP: {args.destino}")
            compactar_arquivos(caminhos_lista, args.destino)
        else:
            print(f"Origem:  {args.origem}")
            print(f"Destino: {args.destino}")
            # Se for uma extração de pasta para ZIP
            if os.path.splitext(args.origem)[1].lower() in (".zip", ".rar", ".7z") and not args.destino.endswith(tuple(MAPA_FORMATOS.keys())):
                print("Operação: Descompactação de Arquivo")
                descompactar_arquivo(args.origem, args.destino)
            else:
                converter_arquivo(args.origem, args.destino, args.delim)
                
        print("\n[SUCESSO] Operação finalizada com êxito!")
    except ConversionError as ce:
        print(f"\n[ERRO DE CONVERSÃO] {str(ce)}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"\n[ERRO INESPERADO] {str(e)}", file=sys.stderr)
        sys.exit(1)
